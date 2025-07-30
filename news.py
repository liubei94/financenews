# news.py (Refactored and Unified Version)

import requests
from bs4 import BeautifulSoup
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import urlparse
from datetime import datetime
import re
import time
import sys

# --- Selenium
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By

# Load environment variables
load_dotenv()
client = OpenAI()

NAVER_CLIENT_ID = os.getenv("NAVER_CLIENT_ID")
NAVER_CLIENT_SECRET = os.getenv("NAVER_CLIENT_SECRET")

# 🔹 HTML → Markdown 변환기
from html_to_markdown import convert_to_markdown

# ✅ 새롭게 정의된 기사 추출 함수
def extract_article(url):
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')

    try:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    except Exception as e:
        print(f"[오류] Chrome 드라이버 설정 실패: {e}")
        return None, None

    try:
        print(f"\n[extract_article] '{url}' 로딩 중...")
        driver.get(url)
        WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(3)
        html_content = driver.page_source
    except Exception as e:
        print(f"[오류] Selenium 로딩 실패: {e}")
        driver.quit()
        return None, None
    finally:
        driver.quit()

    soup = BeautifulSoup(html_content, 'html.parser')
    article_body = None

    selectors = [
        'article', '#article_body', '#dic_area', '#article-view-content-div',
        '#main-content', '#content', '.article_body', '.entry-content', 'main',
    ]
    for selector in selectors:
        found = soup.select_one(selector)
        if found:
            article_body = found
            break
    if not article_body:
        article_body = soup.find('body')
    if not article_body:
        print("[extract_article] 본문 추출 실패")
        return None, None

    title_tag = article_body.find(['h1', 'h2', 'h3']) or soup.find(['h1', 'h2', 'h3']) or soup.find('title')
    page_title = title_tag.get_text(strip=True) if title_tag else "Untitled"
    markdown_content = convert_to_markdown(str(article_body))
    
    return page_title, markdown_content

# ✅ GPT로 키워드 추출
def extract_keywords_with_gpt(title, content):
    prompt = f"""
다음은 뉴스의 제목과 본문입니다. 핵심 키워드 3개를 한글로 추출해주세요.
- 제목에 등장하는 단어나 표현을 우선 고려해 키워드를 선택해주세요.
- 본문 전체를 참고하되, 주제를 잘 대표하는 단어를 뽑아주세요.

제목: {title}
본문: {content}
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "당신은 키워드 추출기입니다. 본문을 가장 잘 나타낼 수 있는 키워드를 추출하세요."},
            {"role": "user", "content": prompt}
        ]
    )
    keywords = response.choices[0].message.content.strip().split('\n')
    cleaned = [re.sub(r'^\d+\.\s*', '', kw).strip() for kw in keywords if kw.strip()]
    return cleaned[:3]

# ✅ 네이버 뉴스 검색
def search_news_naver(keywords, start_date, end_date, display=30):
    query = ' '.join(keywords)
    url = "https://openapi.naver.com/v1/search/news.json"
    headers = {
        "X-Naver-Client-Id": NAVER_CLIENT_ID,
        "X-Naver-Client-Secret": NAVER_CLIENT_SECRET
    }
    params = {"query": query, "display": display, "sort": "date"}
    response = requests.get(url, headers=headers, params=params)
    if response.status_code == 200:
        return response.json().get('items', [])
    else:
        print(f"⚠️ 네이버 API 요청 실패: {response.text}")
        return []

# ✅ 날짜 필터링
def filter_news_by_date(news_items, start_date, end_date):
    start = datetime.strptime(start_date, "%Y-%m-%d").date()
    end = datetime.strptime(end_date, "%Y-%m-%d").date()
    filtered = []
    for item in news_items:
        pub_raw = item.get("pubDate")
        if not pub_raw: continue
        try:
            pub_date = datetime.strptime(pub_raw, "%a, %d %b %Y %H:%M:%S %z").date()
            if start <= pub_date <= end:
                filtered.append(item)
        except ValueError:
            print(f"⚠️ 날짜 파싱 오류: {pub_raw}")
    return filtered

# ✅ GPT로 뉴스 요약
def summarize_news_articles(titles, contents):
    full_text = ""
    for i, (title, content) in enumerate(zip(titles, contents)):
        full_text += f"[{i+1}] {title}\n{content[:800]}\n\n"

    prompt = f"""
당신은 전문 뉴스 분석가입니다. 아래는 여러 뉴스 기사들의 제목과 본문입니다.
이 내용을 **심층 분석 요약** 형식으로 정리해주세요. 요약은 다음 구조를 반드시 따르세요:

---
1. 📌 **핵심 주제 요약** (1~2문장)
2. 📰 **뉴스 요점 정리** (어떤 사건/원인/주요 인물 등)
3. 📊 **비교 또는 이슈 요약 (필요시 표로)**
4. 🧠 **결론 및 시사점** (향후 전망, 독자의 통찰)
---

분석할 뉴스 전체 내용:
{full_text}
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "당신은 정확하고 깊이 있는 뉴스 요약가입니다."}, {"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content.strip()

# ✅ Word 저장
def save_summary_to_word(summary_text, titles, links, news_items, keywords, output_stream, failed_links=None):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)

    section_titles = ["핵심 주제 요약", "뉴스 요점 정리", "비교 또는 이슈 요약", "결론 및 시사점"]
    lines = summary_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue
        if line.startswith('|') and line.endswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                table_data.append(cells)
                i += 1
            if not table_data: continue
            table = doc.add_table(rows=0, cols=len(table_data[0]))
            table.style = 'Table Grid'
            for row_data in table_data:
                row_cells = table.add_row().cells
                for idx, cell_text in enumerate(row_data):
                    row_cells[idx].text = cell_text
            continue
        matched = False
        for title in section_titles:
            if title in line:
                p = doc.add_paragraph()
                run = p.add_run(re.sub(r'^\d+\.\s*📌?\s*', '', line).strip())
                run.bold = True
                run.font.size = Pt(14)
                matched = True
                break
        if matched:
            i += 1; continue
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            else:
                p.add_run(part)
        i += 1

    doc.add_page_break()
    doc.add_paragraph("📎 참고 뉴스 목록", style='Heading 1')
    for idx, (title, link, item) in enumerate(zip(titles, links, news_items), 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. ")
        add_hyperlink(p, link, title)
        origin = extract_news_source(link)
        pubdate = extract_pubdate_from_item(item)
        info = f" ({origin}" + (f", {pubdate}" if pubdate else "") + ")"
        p.add_run(info)

    if failed_links:
        doc.add_paragraph("\n❌ 크롤링 실패 링크", style='Heading 2')
        for link in failed_links:
            doc.add_paragraph(link, style='List Bullet')

    doc.save(output_stream)

# 하이퍼링크 추가 도우미
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'Hyperlink')
    rPr.append(r_style)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def extract_news_source(link):
    netloc = urlparse(link).netloc
    return netloc.replace("www.", "").replace("n.news.", "")

def extract_pubdate_from_item(item):
    if "pubDate" in item:
        try:
            dt = datetime.strptime(item["pubDate"], "%a, %d %b %Y %H:%M:%S %z")
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            return None
    return None

# ✅ 메인 실행 흐름
if __name__ == "__main__":
    test_url = "https://n.news.naver.com/mnews/article/014/0005371160"
    start_date = "2025-07-28"
    end_date = "2025-07-30"

    print("▶️ [1/5] 기사 원문 수집 중...")
    title, content = extract_article(test_url)
    print(f"   - 원문 제목: {title}")

    print("▶️ [2/5] GPT 키워드 추출 중...")
    keywords = extract_keywords_with_gpt(title, content)
    print(f"   - 추출된 키워드: {', '.join(keywords)}")

    print("▶️ [3/5] 네이버 뉴스 검색 및 필터링 중...")
    news_items = search_news_naver(keywords, start_date, end_date)
    filtered_items = filter_news_by_date(news_items, start_date, end_date)
    print(f"   - 검색된 기사 수: {len(filtered_items)}건")

    if filtered_items:
        print("▶️ [4/5] 관련 뉴스 크롤링 및 요약 중...")
        links = [item['link'] for item in filtered_items]
        titles, contents, failed = [], [], []
        for link in links:
            t, c = extract_article(link)
            if t and c:
                titles.append(t)
                contents.append(c)
            else:
                failed.append(link)

        summary = summarize_news_articles(titles, contents)
        print("   - 요약 완료.")

        print("▶️ [5/5] Word 파일 저장 중...")
        output_filename = "news_summary_output.docx"
        save_summary_to_word(summary, titles, links, filtered_items, keywords, output_filename, failed_links=failed)
        print(f"✅ 최종 리포트가 '{output_filename}'으로 저장되었습니다.")
    else:
        print("❌ 날짜에 맞는 뉴스가 없어 리포트를 생성하지 않았습니다.")
