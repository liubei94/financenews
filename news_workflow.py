import requests
from bs4 import BeautifulSoup
from openai import AsyncOpenAI
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from urllib.parse import urlparse
from datetime import datetime
import re
import asyncio
import httpx
from tqdm.asyncio import tqdm
import json

# --- [NEW] crawl4ai and Pydantic imports ---
from pydantic import BaseModel, Field
from crawl4ai import AsyncWebCrawler, CacheMode
from crawl4ai import CrawlerRunConfig, BrowserConfig
from crawl4ai import LLMConfig
from crawl4ai.extraction_strategy import LLMExtractionStrategy
# ---------------------------------------------

# Load environment variables
load_dotenv()

# 비동기 OpenAI 클라이언트 초기화
client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Google Gemini 설정 - 필수 사용
try:
    import google.generativeai as genai
    # Using GOOGLE_API_KEY to align with genai and crawl4ai usage
    google_api_key = os.getenv("GOOGLE_API_KEY")
    if not google_api_key or not google_api_key.strip():
        raise ValueError("GOOGLE_API_KEY 환경변수가 설정되지 않았습니다.")
    genai.configure(api_key=google_api_key)
    print("✅ Gemini API 설정 완료")
except ImportError:
    raise ImportError("google-generativeai 패키지가 설치되지 않았습니다. 'pip install google-generativeai'로 설치해주세요.")
except Exception as e:
    raise Exception(f"Gemini API 설정 실패: {e}")

NAVER_CLIENT_ID = os.getenv("NAVER_CLIENT_ID")
NAVER_CLIENT_SECRET = os.getenv("NAVER_CLIENT_SECRET")


# --- Pydantic model for crawl4ai data structure ---
class NewsArticle(BaseModel):
    title: str = Field(..., description="The main headline or title of the news article.")
    content: str = Field(..., description="The full body text of the news article, excluding ads, comments, and navigation links.")
# ----------------------------------------------------



### 기능 함수들 (Streamlit에서 호출)

async def extract_initial_article_content_async(url: str):
    """
    스크립트 시작 시 기준이 되는 첫 기사를 비동기적으로 가져옵니다.
    crawl4ai를 사용하여 어떤 뉴스 사이트든 안정적으로 제목과 본문을 추출합니다.
    """
    config = CrawlerRunConfig(
        extraction_strategy=LLMExtractionStrategy(
            llm_config=LLMConfig(
                provider="gemini/gemini-1.5-flash",
                api_token=os.getenv("GOOGLE_API_KEY")
            ),
            schema=NewsArticle.model_json_schema(),
            instruction="""Extract the title and the main content of the news article.
            Focus only on the article's body, ignoring comments, related articles, and advertisements.
            Return the result in JSON format based on the provided schema.""",
        ),
        cache_mode=CacheMode.DISABLED
    )
    try:
        async with AsyncWebCrawler(verbose=False) as crawler:
            result = await crawler.arun(url=url, config=config)

        if result.success and result.extracted_content:
            extracted_data = json.loads(result.extracted_content)

            # --- [수정됨] LLM이 리스트를 반환하는 경우에 대한 방어 코드 추가 ---
            article_dict = None
            if isinstance(extracted_data, list) and extracted_data:
                article_dict = extracted_data[0]
            elif isinstance(extracted_data, dict):
                article_dict = extracted_data

            if article_dict:
                title = article_dict.get("title")
                content = article_dict.get("content")
                if not title or not content:
                    raise Exception("crawl4ai가 초기 기사에서 유효한 제목과 본문을 추출하지 못했습니다.")
                print("✅ crawl4ai 초기 기사 추출 성공!")
                return title, content
            # ----------------------------------------------------------
            
            raise Exception("crawl4ai가 초기 기사에서 유효한 데이터를 추출하지 못했습니다.")
        else:
            print(f"❌ crawl4ai 초기 기사 추출 실패: {result.error_message}")
            raise Exception(f"crawl4ai 초기 기사 추출 실패: {result.error_message}")
    except Exception as e:
        print(f"❌ 초기 기사 추출 전체 과정 실패: {e}")
        raise

async def extract_keywords_with_gemini(title, content, max_count=5):
    """Gemini를 사용해 비동기적으로 핵심 키워드를 추출합니다."""
    prompt = f"""
다음은 뉴스의 제목과 본문입니다. 이 기사의 핵심 주제를 가장 잘 나타내는 키워드를 최대 {max_count}개까지 한글로 추출해주세요.
- 각 키워드는 명사 형태로 간결하게 제시해주세요.
- 구분자는 쉼표(,) 또는 줄바꿈을 사용해주세요.

제목: {title}
본문: {content}
"""
    
    def generate_keywords_sync():
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            contents = [prompt]
            generation_config = {"temperature": 0.2}

            response = model.generate_content(
                contents=contents,
                generation_config=generation_config
            )

            if not response.parts:
                print("⚠️ Gemini API가 키워드에 대해 빈 응답을 반환했습니다.")
                return ""
            return response.text.strip()
        except Exception as e:
            print(f"❌ Gemini 키워드 추출 중 오류 발생: {e}")
            return ""

    try:
        keywords_text = await asyncio.to_thread(generate_keywords_sync)
        if not keywords_text:
            raise Exception("Gemini로부터 키워드를 받지 못했습니다.")

        lines = keywords_text.split("\n")
        keywords_list = []
        for line in lines:
            keywords_list.extend([kw.strip() for kw in line.split(",")])

        cleaned_keywords = [
            re.sub(r"^\s*[\d\.\-]+\s*", "", kw).strip()
            for kw in keywords_list
            if kw.strip()
        ]

        return cleaned_keywords[:max_count]
    except Exception as e:
        print(f"❌ 키워드 처리 중 오류 발생: {e}")
        raise

def search_news_naver(keywords, display=50):
    """네이버 API를 통해 관련 뉴스를 검색합니다."""
    query = " ".join(keywords)
    url = "https://openapi.naver.com/v1/search/news.json"
    headers = {
        "X-Naver-Client-Id": NAVER_CLIENT_ID,
        "X-Naver-Client-Secret": NAVER_CLIENT_SECRET,
    }
    params = {"query": query, "display": display, "sort": "date"}
    try:
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        return response.json()["items"]
    except requests.RequestException as e:
        print(f"⚠️ 네이버 API 요청 실패: {e}")
        raise


def filter_news_by_date(news_items, start_date, end_date):
    """검색된 뉴스를 지정된 날짜 범위로 필터링합니다."""
    filtered_items = []
    for item in news_items:
        pub_date_str = item.get("pubDate")
        if not pub_date_str:
            continue
        try:
            pub_date = datetime.strptime(
                pub_date_str, "%a, %d %b %Y %H:%M:%S %z"
            ).date()
            if start_date <= pub_date <= end_date:
                filtered_items.append(item)
        except ValueError:
            continue
    return filtered_items

# [추가] 네이버 뉴스 전용 빠른 스크래퍼 함수
async def extract_naver_article_fast_async(link: str):
    """
    n.news.naver.com 링크에 대해 httpx와 BeautifulSoup를 사용해 빠르게 내용을 추출합니다.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    try:
        async with httpx.AsyncClient() as session:
            response = await session.get(link, headers=headers, timeout=10, follow_redirects=True)
            response.raise_for_status()
            
        soup = BeautifulSoup(response.text, "html.parser")
        
        # 네이버 뉴스 타이틀 추출
        title_tag = soup.find("meta", property="og:title")
        title = title_tag["content"].strip() if title_tag else "제목 없음"
        
        # 네이버 뉴스 본문 선택자 (가장 일반적인 2가지)
        content_area = soup.select_one("article#dic_area, div#newsct_article")
        
        if content_area:
            # 불필요한 요소 제거 (예: 기자 정보, 저작권 문구 등)
            for el in content_area.select("span.byline, div.journalist_area, p.copyright"):
                el.decompose()
            content = content_area.get_text(separator="\n", strip=True)
            # 성공 시 제목과 본문 반환
            if content:
                print(f"✅ 빠른 스크래핑 성공: {link}")
                return title, content

    except Exception as e:
        print(f"⚠️ 빠른 스크래핑 실패: {link}, 오류: {e}")
        pass # 실패 시 아래 crawl4ai가 처리하도록 None 반환
        
    return None, None

async def extract_article_content_async(link: str):
    """
    crawl4ai를 사용하여 비동기적으로 기사 제목과 본문을 추출합니다.
    """
    config = CrawlerRunConfig(
        extraction_strategy=LLMExtractionStrategy(
            llm_config=LLMConfig(
                provider="gemini/gemini-1.5-flash",
                api_token=google_api_key
            ),
            schema=NewsArticle.model_json_schema(),
            instruction="""Extract the title and the main content of the news article.
            Focus only on the article's body, ignoring comments, related articles, and advertisements.
            Return the result in JSON format based on the provided schema.""",
        ),
    )
    try:
        async with AsyncWebCrawler(verbose=False) as crawler:
            result = await crawler.arun(url=link, config=config)
        
        if result.success and result.extracted_content:
            extracted_data = json.loads(result.extracted_content)

            # --- [수정됨] LLM이 리스트를 반환하는 경우에 대한 방어 코드 추가 ---
            article_dict = None
            if isinstance(extracted_data, list) and extracted_data:
                article_dict = extracted_data[0]
            elif isinstance(extracted_data, dict):
                article_dict = extracted_data
            
            if article_dict:
                return article_dict.get("title"), article_dict.get("content")
            # ----------------------------------------------------------

        return None, None
    except Exception:
        return None, None


async def summarize_individual_article_async(title, content):
    prompt = f"""
다음 뉴스 기사의 핵심 내용을 아래 항목에 맞추어 간결하게 요약해줘. 각 항목은 2~3 문장으로 작성해줘.
- **사건/주제**: \n- **주요 인물/기관**: \n- **핵심 주장/내용**: \n- **결과/영향**:
---
제목: {title}\n본문: {content}
"""
    def generate_summary_sync():
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            system_prompt = "당신은 뉴스 분석가입니다. 기사의 핵심만 정확하게 추출하여 요약합니다."
            contents = [system_prompt, prompt]
            generation_config = {"temperature": 0.2}

            response = model.generate_content(
                contents=contents,
                generation_config=generation_config
            )

            if not response.parts:
                print("⚠️ Gemini API가 빈 응답을 반환했습니다.")
                return None
            return response.text.strip()
        except Exception as e:
            print(f"❌ Gemini 개별 요약 중 오류 발생: {e}")
            return None

    try:
        summary = await asyncio.to_thread(generate_summary_sync)
        return summary
    except Exception:
        return None


async def process_article_task(item, semaphore):
    async with semaphore:
        # [수정] 네이버 링크를 우선 사용하도록 순서를 변경합니다.
        link = item.get("link", item.get("originallink"))
        
        title, content = None, None

        # 1. n.news.naver.com 링크인 경우, 먼저 빠른 네이버 전용 스크래퍼를 시도합니다.
        if "n.news.naver.com" in link:
            title, content = await extract_naver_article_fast_async(link)

        # 2. 빠른 스크래핑에 실패했거나, 네이버 링크가 아닐 경우, crawl4ai를 이용한 강력한 폴백(Fallback)을 실행합니다.
        if not title or not content:
            print(f"➡️ crawl4ai 폴백 실행: {link}")
            title, content = await extract_article_content_async(link)
        
        # 두 방식 모두 실패한 경우
        if not title or not content:
            return {"status": "failed", "reason": "크롤링 실패", "link": link}
        
        # 요약 진행
        summary = await summarize_individual_article_async(title, content)
        if not summary:
            return {"status": "failed", "reason": "개별 요약 실패", "link": link}
            
        return {
            "status": "success",
            "title": re.sub("<.*?>", "", item["title"]),
            "link": link,
            "original_item": item,
            "summary": summary,
        }


async def synthesize_final_report(summaries):
    MAX_INPUT_CHARS = 25000 
    
    full_summary_text = ""
    processed_count = 0
    for summary_data in summaries:
        summary_entry = f"### 뉴스 {processed_count + 1}: {summary_data['title']}\n{summary_data['summary']}\n\n---\n\n"
        if len(full_summary_text) + len(summary_entry) > MAX_INPUT_CHARS:
            break
        full_summary_text += summary_entry
        processed_count += 1

    if processed_count < len(summaries):
        remaining_count = len(summaries) - processed_count
        full_summary_text += f"\n... 외 {remaining_count}개의 관련 기사가 있으나, 안정적인 분석을 위해 일부만 사용합니다.\n"

    system_prompt = """
당신은 정치/경제/산업 분야의 최고 수준의 전문 분석가입니다. 
여러 뉴스 기사의 핵심 요약본들을 바탕으로, 회사 CFO나 CEO가 의사결정을 위해 참고할 심층 분석 보고서를 작성합니다.

# 지시사항
아래의 구조와 세부 지침을 반드시 준수하여 보고서를 작성해주세요.
전체적인 분량은 2000자 내외로, Key Developments와 Comparative Analysis의 비중을 각각 40% 수준을 유지합니다.
단, 최종 결과물을 출력할 때 목록은 번호(1., 2.)를 사용하지 말고, 오직 불렛 포인트('*')로만 통일해야 합니다.

---
## 보고서 구조 및 세부 지침
1.  **📌 Executive Summary (핵심 요약)**
    * 전체 상황을 1~2 문장으로 요약합니다.
2.  **📰 Key Developments (주요 동향 및 사실 분석)**
    * 어떤 사건/행동이 있었는지 종합적으로 설명합니다.
    * 공통적으로 드러나는 원인과 배경은 무엇입니까?
    * 핵심적인 이해관계자(인물, 기업, 기관)는 누구이며, 그들의 입장은 무엇입니까?
3.  **📊 Comparative Analysis (비교 분석 및 이슈 심층 탐구)**
    * 기사들 간의 관점 차이나 상충되는 정보가 있다면 비교 분석합니다.
    * 수치, 데이터, 정책 변화 등 중요한 포인트를 표(Table) 형식으로 정리하여 시각적 이해를 돕습니다. (필요시)
4.  **🧠 Conclusion & Strategic Implications (결론 및 전략적 시사점)**
    * Key Developments와 Comparative Analysis의 내용에서 벗어나지 않도록 주의합니다.
    * 이러한 흐름이 향후 시장/산업/정책에 미칠 영향은 무엇입니까?
    * 우리 조직이 주의 깊게 관찰해야 할 리스크와 기회 요인은 무엇입니까? 
    * 독자가 얻어야 할 최종적인 통찰(Insight)을 제시합니다.
주의사항 : 날짜, 수신자, 참조 등 보고서 헤더 정보는 생성하지 않고 본문 내용만 작성합니다.
"""
    user_prompt = f"아래의 뉴스 요약본들을 바탕으로 분석 보고서를 작성해주세요.\n\n---## 요약본 시작 ##---\n\n{full_summary_text}"

    def generate_content_sync():
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            generation_config = {"temperature": 0.2} 
            response = model.generate_content(
                contents=[system_prompt, user_prompt],
                generation_config=generation_config
            )
            if not response.parts:
                raise ValueError(f"Gemini API가 빈 응답을 반환했습니다. (finish_reason: {response.candidates[0].finish_reason.name})")
            return response.text.strip()
        except Exception as e:
            raise Exception(f"최종 보고서 생성 중 Gemini API 오류: {e}")

    return await asyncio.to_thread(generate_content_sync)


async def run_analysis_and_synthesis_async(filtered_items, progress_callback=None):
    semaphore = asyncio.Semaphore(10)
    successful_results = []
    failed_results = []
    total_items = len(filtered_items)

    tasks = [process_article_task(item, semaphore) for item in filtered_items]
    for i, future in enumerate(asyncio.as_completed(tasks)):
        result = await future
        if result and result["status"] == "success":
            successful_results.append(result)
        else:
            failed_results.append(result or {"status": "failed", "reason": "알 수 없는 오류", "link": ""})

        if progress_callback:
            progress_callback(i + 1, total_items, f"📰 기사 요약 중... ({i + 1}/{total_items})")

    if not successful_results:
        return None, [], []

    if progress_callback:
        progress_callback(total_items, total_items, "✅ 분석 완료! 최종 보고서를 종합합니다...")

    final_report = await synthesize_final_report(successful_results)

    return final_report, successful_results, failed_results


# --- Word 저장 로직 ---
def save_summary_to_word(summary_text, successful_results, output_stream):
    """분석 결과를 서식이 적용된 Word 문서로 저장합니다."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(5)
    
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("AI 뉴스 분석 리포트")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_p.alignment = 1

    today_str = datetime.now().strftime('%Y년 %m월 %d일')
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"작성일: {today_str}")
    date_run.font.size = Pt(11)
    date_p.alignment = 2

    doc.add_paragraph("---")

    lines = summary_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line or line == "---":
            continue

        p = None

        if line.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(line.replace("### ", "")).bold = True
            p.runs[0].font.size = Pt(12)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.add_run(line.replace("## ", "")).bold = True
            p.runs[0].font.size = Pt(14)
        elif line.startswith("# ") or line.startswith("📌") or line.startswith("📰") or line.startswith("📊") or line.startswith("🧠"):
            p = doc.add_paragraph()
            p.add_run(line.lstrip("# 📌📰📊🧠").strip()).bold = True
            p.runs[0].font.size = Pt(16)
        elif line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(20) 
            clean_line = line.replace("* ", "").replace("**", "")
            p.add_run(clean_line)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("📎 참고 뉴스 목록")
    run.bold = True
    run.font.size = Pt(16)

    for idx, result in enumerate(successful_results, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. ")
        add_hyperlink(p, result["link"], result["title"])
        origin = extract_news_source(result["link"])
        pubdate = extract_pubdate_from_item(result["original_item"])
        info = f" ({origin}" + (f", {pubdate}" if pubdate else "") + ")"
        p.add_run(info)

    doc.save(output_stream)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def extract_news_source(link):
    try:
        netloc = urlparse(link).netloc
        parts = netloc.replace("www.", "").split(".")
        domain = parts[-2] if len(parts) > 1 else parts[0]
        source_map = {
            "chosun": "조선일보",
            "donga": "동아일보",
            "mk": "매일경제",
            "joongang": "중앙일보",
            "hani": "한겨레",
            "yna": "연합뉴스",
            "inews24": "아이뉴스24",
            "fnnews": "파이낸셜뉴스",
            "naver": "네이버뉴스",
        }
        return source_map.get(domain, domain)
    except:
        return "알 수 없는 출처"


def extract_pubdate_from_item(item):
    if "pubDate" in item:
        try:
            dt = datetime.strptime(item["pubDate"], "%a, %d %b %Y %H:%M:%S %z")
            return dt.strftime("%Y-%m-%d")
        except:
            return None
    return None
